"""
Advanced Embedding Service with Hybrid Retrieval - MongoDB + Elasticsearch Version
Handles document embedding, Elasticsearch BM25 indexing, and similarity search with group-based filtering
"""

import os
import json
import time
import logging
from typing import List, Dict, Optional
from datetime import datetime, timedelta
import redis
import numpy as np
from sentence_transformers import SentenceTransformer
import chromadb
from  pathlib import Path
from chromadb.config import Settings
from elasticsearch import Elasticsearch
import PyPDF2
import docx

# Import advanced RAG components
from rag_utils import (
    Document, AdvancedRAGPipeline,
    initialize_chroma_client, get_mongodb_connection, get_elasticsearch_connection, HybridRetriever
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Configuration
REDIS_URL = f"redis://:{os.environ.get('REDIS_PASSWORD', 'password')}@{os.environ.get('REDIS_HOST', 'localhost')}:{os.environ.get('REDIS_PORT', '6379')}/{os.environ.get('REDIS_DB', '0')}"
CHROMA_HOST = os.environ.get('CHROMA_HOST', 'localhost')
CHROMA_PORT = int(os.environ.get('CHROMA_PORT', '8000'))
CHROMA_PERSISTENT = os.environ.get('CHROMA_PERSISTENT', 'false').lower() == 'true'
EMBEDDING_MODEL = os.environ.get('EMBEDDING_MODEL', 'all-MiniLM-L6-v2')
RERANKER_MODEL = os.environ.get('RERANKER_MODEL', 'cross-encoder/ms-marco-MiniLM-L-6-v2')

# MongoDB configuration
MONGO_HOST = os.environ.get('MONGODB_HOST', 'localhost')
MONGO_PORT = os.environ.get('MONGO_INITDB_PORT', '27017')
MONGO_USERNAME = os.environ.get('MONGO_INITDB_ROOT_USERNAME', 'admin')
MONGO_PASSWORD = os.environ.get('MONGO_INITDB_ROOT_PASSWORD', 'password')
MONGO_DATABASE = os.environ.get('MONGO_INITDB_DATABASE', 'rag_system')

MONGO_URI = f"mongodb://{MONGO_USERNAME}:{MONGO_PASSWORD}@{MONGO_HOST}:{MONGO_PORT}"

# Elasticsearch configuration
ELASTICSEARCH_HOST = os.environ.get('ELASTICSEARCH_HOST', 'localhost')
ELASTICSEARCH_PORT = int(os.environ.get('ELASTICSEARCH_PORT', '9200'))
ELASTICSEARCH_USERNAME = os.environ.get('ELASTICSEARCH_USERNAME', None)
ELASTICSEARCH_PASSWORD = os.environ.get('ELASTICSEARCH_PASSWORD', None)
ELASTICSEARCH_INDEX = os.environ.get('cluster.name', 'rag-elastic-cluster')

# Global variables for services
redis_client = None
embedding_model = None
chroma_client = None
es_client = None
mongo_client = None
db = None
rag_pipeline = None


def initialize_services():
    """Initialize all services with proper error handling"""
    global redis_client, embedding_model, chroma_client, es_client, mongo_client, db, rag_pipeline
    
    logger.info("Initializing services...")
    
    # Initialize Redis
    try:
        redis_client = redis.from_url(REDIS_URL, decode_responses=True)
        redis_client.ping()
        logger.info("Redis connected successfully")
    except Exception as e:
        logger.error(f"Redis connection failed: {e}")
        raise
    
    # Initialize embedding model
    try:
        embedding_model = SentenceTransformer(EMBEDDING_MODEL)
        logger.info(f"Embedding model {EMBEDDING_MODEL} loaded successfully")
    except Exception as e:
        logger.error(f"Failed to load embedding model: {e}")
        raise
    
    # Initialize ChromaDB
    try:
        if CHROMA_PERSISTENT:
            chroma_client = initialize_chroma_client(use_persistent=True)
        else:
            chroma_client = initialize_chroma_client(
                host=CHROMA_HOST, 
                port=CHROMA_PORT, 
                use_persistent=False
            )
        logger.info("ChromaDB initialized successfully")
    except Exception as e:
        logger.error(f"ChromaDB initialization failed: {e}")
        # Fallback to in-memory client
        chroma_client = chromadb.Client(settings=Settings(anonymized_telemetry=False))
        logger.warning("Using in-memory ChromaDB client")
    
    # Initialize Elasticsearch
    try:
        es_client = get_elasticsearch_connection(
            host=ELASTICSEARCH_HOST,
            port=ELASTICSEARCH_PORT,
            username=ELASTICSEARCH_USERNAME,
            password=ELASTICSEARCH_PASSWORD
        )
        logger.info("Elasticsearch connected successfully")
    except Exception as e:
        logger.error(f"Elasticsearch connection failed: {e}")
        raise
    
    # Initialize MongoDB
    try:
        mongo_client, db = get_mongodb_connection(MONGO_URI, MONGO_DATABASE)
        logger.info("MongoDB connected successfully")
    except Exception as e:
        logger.error(f"MongoDB connection failed: {e}")
        raise
    
    # Initialize RAG pipeline
    try:
        rag_pipeline = AdvancedRAGPipeline(
            mongo_uri=MONGO_URI,
            database_name=MONGO_DATABASE,
            redis_client=redis_client,
            chroma_client=chroma_client,
            es_client=es_client
        )
        logger.info("Advanced RAG pipeline initialized successfully")
    except Exception as e:
        logger.error(f"RAG pipeline initialization failed: {e}")
        raise
    
    logger.info("All services initialized successfully")


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from different file types"""
    try:
        logger.info(f"Extracting text from {file_type} file: {file_path}")
        
        if file_type == 'pdf':
            text = ""
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for i, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            # Add page markers for better chunking
                            text += f"\n\n[Page {i+1}]\n{page_text}"
                logger.info(f"Extracted {len(text)} characters from PDF")
                return text
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
                return ""
        
        elif file_type in ['doc', 'docx']:
            text = ""
            try:
                doc = docx.Document(file_path)
                
                # Extract from paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n\n"
                
                # Extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            text += row_text + "\n"
                    text += "\n"
                
                logger.info(f"Extracted {len(text)} characters from DOCX")
                return text
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                return ""
        
        elif file_type == 'txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                logger.info(f"Extracted {len(text)} characters from TXT")
                return text
            except UnicodeDecodeError:
                # Try different encodings
                for encoding in ['latin1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as file:
                            text = file.read()
                        logger.info(f"Extracted {len(text)} characters from TXT using {encoding}")
                        return text
                    except:
                        continue
                logger.error("Failed to decode text file with any encoding")
                return ""
        
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return ""
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return ""


def get_file_type_from_filename(filename: str) -> str:
    """
    Determine file type from filename extension
    Returns the file type or raises ValueError for unsupported types
    """
    if not filename:
        raise ValueError("Filename cannot be empty")
    
    # Get file extension (convert to lowercase for case-insensitive comparison)
    file_extension = os.path.splitext(filename.lower())[1]
    
    # Define supported file types
    supported_extensions = {
        '.pdf': 'pdf',
        '.txt': 'txt',
        '.doc': 'doc',
        '.docx': 'docx'
    }
    
    if file_extension not in supported_extensions:
        supported_types = ', '.join(supported_extensions.keys())
        raise ValueError(f"Unsupported file type '{file_extension}'. Supported types: {supported_types}")
    
    return supported_extensions[file_extension]

def process_uploaded_file(task: Dict):
    """Process an uploaded file with advanced chunking and indexing"""
    try:
        filename = task['filename']
        file_id = task['file_id']
        file_path = task['file_path']
        group_id = task.get('group_id', 'general')  # Get group_id from task
        
        logger.info(f"Processing file {file_id}: {file_path} for group {group_id}")
        
        # Check if file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type from filename
        try:
            file_type = get_file_type_from_filename(filename)
            logger.info(f"Detected file type: {file_type} for file: {filename}")
        except ValueError as e:
            logger.error(f"File type determination failed for {filename}: {e}")
            raise ValueError(f"Unsupported file type: {e}")
        
        # Extract text from file
        text = extract_text_from_file(file_path, file_type)
        
        if not text or len(text.strip()) < 10:
            raise ValueError("No meaningful text extracted from file")
        
        logger.info(f"Extracted {len(text)} characters from file {file_id}")
        
        # Create document with filename prominently in metadata
        doc = Document(
            id=file_id,
            content=text,
            metadata={
                'filename': filename,  # Make sure filename is first in metadata
                'file_id': file_id,
                'file_type': file_type,
                'group_id': group_id,
                'source': 'user_upload',
                'upload_date': datetime.utcnow().isoformat(),
                'file_size': len(text),
                'processed_by': 'advanced_rag_pipeline'
            }
        )
        
        # Process with advanced pipeline (group-aware)
        chunks_count = rag_pipeline.process_document(doc, group_id)
        
        # Update file status in MongoDB
        db.file_uploads.update_one(
            {'_id': file_id},
            {
                '$set': {
                    'upload_status': 'completed',
                    'chunks_count': chunks_count,
                    'processed_at': datetime.utcnow(),
                    'text_length': len(text),
                    'file_type': file_type,
                    'filename': filename,  # Store filename in file_uploads collection too
                    'group_id': group_id  # Store group_id
                }
            }
        )
            
        logger.info(f"Successfully processed file {file_id} ({filename}) with {chunks_count} chunks for group {group_id}")
        
        # Store result in Redis
        result = {
            'file_id': file_id,
            'filename': filename,  # Include filename in result
            'status': 'completed',
            'chunks_count': chunks_count,
            'text_length': len(text),
            'file_type': file_type,
            'group_id': group_id,
            'timestamp': datetime.utcnow().isoformat()
        }
        
        try:
            redis_client.setex(
                f"file_processing_result:{file_id}",
                300,  # 5 minutes TTL
                json.dumps(result)
            )
        except:
            logger.warning("Failed to store result in Redis")
        
    except Exception as e:
        logger.error(f"Error processing file {file_id} ({filename}): {e}")
        
        # Update file status to failed
        try:
            db.file_uploads.update_one(
                {'_id': file_id},
                {
                    '$set': {
                        'upload_status': 'failed',
                        'error_message': str(e),
                        'processed_at': datetime.utcnow(),
                        'filename': filename  # Store filename even on failure
                    }
                }
            )
        except Exception as db_error:
            logger.error(f"Failed to update file status in MongoDB: {db_error}")
        
        # Store error result in Redis
        result = {
            'file_id': file_id,
            'filename': filename,  # Include filename in error result
            'status': 'failed',
            'error': str(e),
            'timestamp': datetime.utcnow().isoformat()
        }
        
        try:
            redis_client.setex(
                f"file_processing_result:{file_id}",
                300,
                json.dumps(result)
            )
        except:
            logger.warning("Failed to store error result in Redis")

def process_chat_task(task: Dict):
    """Process chat retrieval task with group-based filtering"""
    try:
        start_time = time.time()
        
        # Get group_id from task (defaults to 'general' if not specified)
        group_id = task.get('group_id', 'general')
        
        logger.info(f"Processing chat query: '{task['message'][:100]}...' for group: {group_id}")
        
        # Retrieve using advanced pipeline with group filtering
        results = rag_pipeline.retrieve(
            query=task['message'],
            k=5,
            group_id=group_id
        )
        
        # Format results for response
        documents = []
        for result in results:
            # Extract filename from metadata if available
            filename = result.get('metadata', {}).get('filename', 'Unknown')
            
            documents.append({
                'id': result['chunk_id'],
                'content': result['content'],
                'metadata': {
                    **result['metadata'],
                    'filename': filename  # Ensure filename is prominently available
                },
                'filename': filename,  # Also add filename at top level for easy access
                'relevance_score': result['final_score'],
                'scores': {
                    'initial': result['initial_score'],
                    'rerank': result['rerank_score'],
                    'final': result['final_score']
                }
            })
        
        retrieval_time = int((time.time() - start_time) * 1000)
        
        # Store result in Redis
        result = {
            'task_id': task['id'],
            'documents': documents,
            'group_id': group_id,
            'retrieval_time_ms': retrieval_time,
            'timestamp': datetime.utcnow().isoformat()
        }
        
        try:
            redis_client.setex(
                f"retrieval_result:{task['id']}", 
                60,
                json.dumps(result)
            )
        except:
            logger.warning("Failed to store retrieval result in Redis")
        
        logger.info(f"Retrieved {len(documents)} documents in {retrieval_time}ms for group {group_id}")
        
    except Exception as e:
        logger.error(f"Chat task processing failed: {e}")
        # Store error result
        error_result = {
            'task_id': task['id'],
            'error': str(e),
            'group_id': task.get('group_id', 'general'),
            'timestamp': datetime.utcnow().isoformat()
        }
        try:
            redis_client.setex(
                f"retrieval_result:{task['id']}", 
                60,
                json.dumps(error_result)
            )
        except:
            logger.warning("Failed to store error result in Redis")

def process_group_deletion_task(task: Dict):
    """Process group deletion task"""
    try:
        group_id = task.get('group_id')
        if not group_id or group_id == 'general':
            logger.warning("Invalid or general group_id for deletion")
            return
        
        logger.info(f"Processing group deletion for group: {group_id}")
        

        # Delete all data for the group
        deletion_counts = rag_pipeline.delete_group_data(group_id)
        
        result = {
            'task_id': task['id'],
            'group_id': group_id,
            'deletion_counts': deletion_counts,
            'timestamp': datetime.utcnow().isoformat()
        }
        
        try:
            redis_client.setex(
                f"group_deletion_result:{task['id']}", 
                60,
                json.dumps(result)
            )
        except:
            logger.warning("Failed to store group deletion result in Redis")
        
        logger.info(f"Completed group deletion for {group_id}: {deletion_counts}")
        
    except Exception as e:
        logger.error(f"Group deletion task failed: {e}")





def process_embedding_tasks():
    """
    Main loop to process embedding/retrieval tasks from Redis
    """
    logger.info("Advanced Embedding service started with MongoDB, ChromaDB, and Elasticsearch")
    
    # Task type handlers
    task_handlers = {
        'chat': process_chat_task,
        'process_file': process_uploaded_file,
        'delete_group': process_group_deletion_task,
    }
    
    while True:
        try:
            # Check for multiple task types
            task_data = redis_client.brpop([
                'embedding_tasks', 
                'file_processing_tasks',
                'group_management_tasks'
            ], timeout=1)
            
            if task_data:
                queue_name, task_json = task_data
                task = json.loads(task_json)
                
                logger.info(f"Processing task from {queue_name}: {task.get('id', 'unknown')} - Type: {task.get('type', 'unknown')}")
                
                # Route task to appropriate handler
                task_type = task.get('type')
                if task_type in task_handlers:
                    task_handlers[task_type](task)
                else:
                    logger.warning(f"Unknown task type: {task_type}")
                
        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in task: {e}")
        except Exception as e:
            logger.error(f"Error processing task: {e}")
            time.sleep(1)
        except KeyboardInterrupt:
            logger.info("Shutting down service...")
            break



def check_chroma_health() -> bool:
    """Check ChromaDB health"""
    try:
        health_status = rag_pipeline.retriever.health_check()
        return health_status.get('chromadb', False)
    except:
        return False


def check_elasticsearch_health() -> bool:
    """Check Elasticsearch health"""
    try:
        health_status = rag_pipeline.retriever.health_check()
        return health_status.get('elasticsearch', False)
    except:
        return False


def check_redis_health() -> bool:
    """Check Redis health"""
    try:
        redis_client.ping()
        return True
    except:
        return False


def check_mongodb_health() -> bool:
    """Check MongoDB health"""
    try:
        db.command('ping')
        return True
    except:
        return False


def check_embedding_model_health() -> bool:
    """Check embedding model health"""
    try:
        test_embedding = embedding_model.encode("test")
        return len(test_embedding) > 0
    except:
        return False


def get_total_indexed_documents(group_id: Optional[str] = None) -> int:
    """Get total number of indexed documents, optionally filtered by group"""
    try:
        query_filter = {}
        if group_id and group_id != "general":
            query_filter['group_id'] = group_id
        return db.indexed_documents.count_documents(query_filter)
    except:
        return 0



def get_elasticsearch_index_stats() -> Dict:
    """Get Elasticsearch index statistics"""
    try:
        stats = es_client.indices.stats(index=ELASTICSEARCH_INDEX)
        index_stats = stats.get('indices', {}).get(ELASTICSEARCH_INDEX, {})
        
        return {
            'document_count': index_stats.get('total', {}).get('docs', {}).get('count', 0),
            'store_size_bytes': index_stats.get('total', {}).get('store', {}).get('size_in_bytes', 0),
            'indexing_total': index_stats.get('total', {}).get('indexing', {}).get('index_total', 0),
            'search_query_total': index_stats.get('total', {}).get('search', {}).get('query_total', 0)
        }
    except:
        return {}


def health_check() -> Dict:
    """Comprehensive health check"""
    return {
        'status': 'healthy',
        'timestamp': datetime.utcnow().isoformat(),
        'services': {
            'chromadb': check_chroma_health(),
            'elasticsearch': check_elasticsearch_health(),
            'redis': check_redis_health(),
            'mongodb': check_mongodb_health(),
            'embedding_model': check_embedding_model_health()
        },
        'metrics': {
            'total_indexed_documents': get_total_indexed_documents(),
            'elasticsearch_stats': get_elasticsearch_index_stats()
        },
        'configuration': {
            'embedding_model': EMBEDDING_MODEL,
            'reranker_model': RERANKER_MODEL,
            'chroma_persistent': CHROMA_PERSISTENT,
            'mongo_database': MONGO_DATABASE,
            'elasticsearch_host': ELASTICSEARCH_HOST,
            'elasticsearch_port': ELASTICSEARCH_PORT,
            'elasticsearch_index': ELASTICSEARCH_INDEX
        }
    }


def cleanup_services():
    """Cleanup services on shutdown"""
    global mongo_client, redis_client, es_client
    
    logger.info("Cleaning up services...")
    
    try:
        if mongo_client:
            mongo_client.close()
            logger.info("MongoDB connection closed")
    except:
        pass
    
    try:
        if redis_client and hasattr(redis_client, 'close'):
            redis_client.close()
            logger.info("Redis connection closed")
    except:
        pass
    
    try:
        if es_client and hasattr(es_client, 'close'):
            es_client.close()
            logger.info("Elasticsearch connection closed")
    except:
        pass


if __name__ == "__main__":
    try:
        status_file = '/app_data/logs/embedding_service_status.txt'
        if Path(status_file).exists():
            os.remove(status_file)
            
        # Initialize all services
        initialize_services()
        
        # Perform health check
        health_status = health_check()
        logger.info(f"Health check: {health_status}")
        
        if all(health_status['services'].values()):
            try:
                with open(status_file, 'w') as f:
                    f.write(f"SUCCESS: Embedding service healthy at {health_status['timestamp']}\n")
                logger.info(f"Success status file created at {status_file}")
            except Exception as e:
                logger.error(f"Failed to write status file: {e}")
                
        # Start the main processing loop
        process_embedding_tasks()
        
    except KeyboardInterrupt:
        logger.info("Service interrupted by user")
    except Exception as e:
        logger.error(f"Service failed to start: {e}")
        raise
    finally:
        cleanup_services()
        logger.info("Service shutdown complete")